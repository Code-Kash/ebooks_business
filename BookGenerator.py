import sys
import os
import openai
import glob
import re
import docx
import argparse

class BookGenerator:

    def __init__(self, topic, model = "text-davinci-003", 
    confirm = False,
    use_existing_outline = False,
    dry_run = False):
        self.topic = topic
        self.min_chapters = 8
        self.max_chapters = 13
        self.min_examples = 0
        self.max_examples = 2
        self.model = model
        self.confirm = confirm
        self.use_existing_outline = use_existing_outline
        self.dry_run = dry_run
        self.outline_text = ""
        self.short_outline_text = ""
        self.input_file_name = None
        
    def _select_outline_file(self):
        existing_outline_files = glob.glob(os.path.join("outlines", f"{self.topic}_outline*"))
        if len(existing_outline_files)> 0:
            if self.use_existing_outline:
                return max(existing_outline_files)

            if self.confirm:
                print(f"The following outline exist for topic {self.topic}.")
                print("\n".join(list([f"\t[{index}] {outline_f}" 
                                        for index, outline_f in enumerate(existing_outline_files)])))
                print("\t[n] Create new outline")

            while True:
                try:
                    if self.confirm:
                        index = input("Select a valid index from below options: ")
                    else:
                        index = "n"
                    if index.lower().strip() == "n":
                        next_index = int(re.search(".(\d+).txt", max(existing_outline_files)).group(1)) + 1
                        return os.path.join("outlines", f"{self.topic}_outline.{next_index:05.0f}.txt")
                    elif int(index) >= 0 and int(index) <= len(existing_outline_files):
                        return existing_outline_files[int(index)]
                    else:
                        print("Invalid choice given!")
                except KeyboardInterrupt:
                    print()
                    sys.exit(0)
                except:
                    pass
    
        return os.path.join("outlines", f"{self.topic}_outline.00001.txt")
    def _init(self):
        selected_outline_file = self._select_outline_file()
        self.input_file_name = os.path.basename(selected_outline_file)
        if selected_outline_file is not None and os.path.isfile(selected_outline_file):
            print(f"reading from file {selected_outline_file}...")
            self.outline_text = open(selected_outline_file, "r").read()
            return
    
        outline_prompt = f"Write me a book outline on {self.topic}. Use your best judgement to determine the number of chapters for the book (choose 8, 9, 10, 11, 12 or 13). " +\
        "Each chapter has 3 topics. Each topic has 2 subtopics.  Chapters are counted with integers. " +\
        "Topics are counted with integers under Chapters with the prefix Topic:. Subtopics are prefixed with only bullet points under topics." +\
        "Propose a title for the book with Book Title prefix." 

        #print(f"PROMPT: {outline_prompt}")
        self.outline_text = self._execute_prompt(outline_prompt)

        print(f"OUTLINE{'-'*10}\n{self.outline_text}")
        print(f"END OUTLINE{'-'*10}")
        if self.confirm:
            while True:
                answer = input("Do you want to continue with this outline? [Y(es)/N(o)/R(edo)]")
                if answer.lower() == "n":
                    sys.exit(0)
                elif answer.lower() == "r":
                    self._init()
                    return
                elif answer.lower() == "y":
                    break
                else:
                    print("Invalid choice, please retry!\n\n")
        print(f"writing to file {selected_outline_file}")
        with open(selected_outline_file, "w") as fout:
            fout.write(f"{self.topic}\n")
            fout.write(self.outline_text)
    
    def _execute_prompt(self, prompt):
        print(f"PROMPT: {prompt}")
        if self.dry_run:
            return ""
        request = openai.Completion.create(
            model=self.model,
            prompt=prompt,
            temperature=0.7,
            max_tokens=4097 - len(prompt),
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0
        )
        return request["choices"][0]["text"]

    def _deserialize_outline_text(self):
        book = dict()
        current_chapter = None
        #current_topic = None
        first_line = True
        for line in self.outline_text.split("\n"):
            if first_line:
                first_line = False
                continue
            if line.lower().startswith("book title"):
                book["title"] = line.lower().replace("book title:", "").strip().title()
                book["Chapters"] = {}
            elif line.lower().startswith("chapter"):
                current_chapter = line.split(":")[0].strip()
                book["Chapters"][current_chapter] = dict()
                book["Chapters"][current_chapter]["name"] = line.split(":")[1].strip()
                book["Chapters"][current_chapter]["topics"] = {}
                book["Chapters"][current_chapter]["main_topics"] = []
            elif line.lower().strip().startswith("topic"):
                current_topic = line.split(":")[0].strip()
                book["Chapters"][current_chapter]["topics"][current_topic] = dict()
                book["Chapters"][current_chapter]["topics"][current_topic]["name"] = line.split(":")[1].strip()
                book["Chapters"][current_chapter]["topics"][current_topic]["subtopics"] = []
                book["Chapters"][current_chapter]["main_topics"].append(re.sub(r'[^A-Za-z0-9 ]+', '', line.split(":")[1].strip()))
            elif line.strip() != "":
                book["Chapters"][current_chapter]["topics"][current_topic]["subtopics"].append(re.sub(r'[^A-Za-z0-9 ]+', '', line.strip()))
        return book

    def generate(self):
        self._init()
        # Parse command line arguments
        parser = argparse.ArgumentParser()
        parser.add_argument('-g', '--glossary', action='store_true')
        parser.add_argument('-e', '--expanded', action='store_true')
        args = parser.parse_args()
        generate_glossary = args.glossary
        expanded = args.expanded

        book = self._deserialize_outline_text()
        short_outline = ""
        for line in self.outline_text.splitlines():
            if line.lower().strip().startswith("chapter") or line.lower().strip().startswith("topic") or line.lower().strip().startswith("book"):
                 short_outline += line + "\n"

        mydoc = docx.Document()
        mydoc.add_heading(book["title"], level=1)
        # Generate book Introduction
        intro_prompt = f"The following is a book outline for an ebook:\n {short_outline}\n" +\
            "Write a detailed 2000 word Introduction for the book."
        mydoc.add_heading("Introduction", level=1)
        mydoc.add_paragraph(self._execute_prompt(intro_prompt))
        # Generate Introduction for each Chapter
        for chapter_order, chapter in book["Chapters"].items():
            chapter_intro_prompt = f"The following is a book outline for an ebook:{short_outline}\n" +\
                f"Write a 2000 word Introduction for {chapter_order}: {chapter['name']}.\n " +\
                f"The Chapter's main topics are: {', '.join(chapter['main_topics'])}"
            mydoc.add_heading(f"{chapter_order}: {chapter['name']}", level=1)
            mydoc.add_paragraph(self._execute_prompt(chapter_intro_prompt))
            # Generate Topic Content
            for topic_order, topic in chapter["topics"].items():
                mydoc.add_heading(topic['name'], level=2)
                if expanded :
                    # Generate paragraphs for subtopics
                    for subtopic in topic['subtopics']:
                        chapter_topic_prompt = f"The following is a book outline for an ebook:{short_outline}\n" +\
                            f"Write a 2000 words section within {chapter_order}: {chapter['name']}. " +\
                            f"The section's topic is: {subtopic}. Choose between 0, 1, and 2 examples to append the end of this section. " +\
                            f"Do not prefix paragraphs with titles."    
                        mydoc.add_paragraph(self._execute_prompt(chapter_topic_prompt))
                else:
                    # Generate paragraphs for topics
                    chapter_topic_prompt = f"The following is a book outline for an ebook:{short_outline}\n" +\
                    f"Write a 2000 words section within {chapter_order}: {chapter['name']}. " +\
                    f"The section's main topic is: {topic['name']}. Choose between 0, 1, and 2 examples to append the end of this section. "
                    mydoc.add_paragraph(self._execute_prompt(chapter_topic_prompt))
        
        # Generate Conclusion
        conclusion_prompt = f"The following is a book outline for an ebook:\n {short_outline}\n" +\
            "As a professional Author, write a detailed 2000 words conclusion for the book."
        mydoc.add_heading("Conclusion", level=1)
        mydoc.add_paragraph(self._execute_prompt(conclusion_prompt))

        if generate_glossary:
            # Generate glossary 
            glossary_prompt = f"The following is a book outline for an ebook:\n {short_outline}\n" +\
                "As a professional Author, write a detailed glossary for the book."
            mydoc.add_heading("Glossary", level=1)
            mydoc.add_paragraph(self._execute_prompt(glossary_prompt))
        # Generate Preface
        # preface_prompt = f"The following is a book outline for an ebook:\n {short_outline}\n" +\
        #     "As a professional Author, write a 2000 words detailed preface for the book."
        # mydoc.add_heading("Preface", level=1)
        # mydoc.add_paragraph(self._execute_prompt(preface_prompt))
        
        output_fname = self.input_file_name.replace("outline", "fullbook").replace(".txt", ".docx")
        mydoc.save(os.path.join("fullbooks", output_fname))
